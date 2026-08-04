from pathlib import Path

from docx import Document

from writersroom.importers.docx_importer import (
    DocxImporter,
)


def main():
    print("Testing DocxImporter...")

    test_file = Path(
        "tests/story.docx"
    )

    document = Document()

    document.add_heading(
        "Story"
    )

    document.add_paragraph(
        "A protagonist should pursue a concrete objective."
    )

    document.save(
        test_file
    )

    importer = DocxImporter()

    result = importer.import_document(
        str(test_file)
    )

    assert (
        result.filename
        == "story.docx"
    )

    assert (
        result.metadata["format"]
        == "docx"
    )

    assert (
        "Story"
        in result.text
    )

    assert (
        "A protagonist should pursue"
        in result.text
    )

    test_file.unlink()

    print(
        "DocxImporter tests passed."
    )


if __name__ == "__main__":
    main()